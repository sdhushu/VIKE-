from tkinter import Tk
import time
import tkinter as tk
from docx import Document
from pygui.function_cc.zhifuSpider.zhifubao import ZhiFuBao
from pygui.Spider.taobaoSpider.taobao2 import taobao_go
from pygui.Spider.bilbilSpider.bilbilSpider import bilibili_go
from pygui.Spider.musicSpider.musicSpider import go
from pygui.function_cc.studySpider.studySpider import Login
from pygui.Spider.qqemailSpider.qq_emailSpiser import emailgo
from pygui.Spider.tuniuSpider.tuniuSpiser import tuniugo
from pygui.Spider.baiduSpider.baiduSpider import baidu_go
from pygui.Spider.jdSpiser.JDSpider import jdgo
from pygui.Spider.cloudcc.cloud import cloudcc
from pygui.Spider.vs.imgchange import ResizeImage
from pygui.Spider.vs.vs import vsgo
import random
from tkinter.messagebox import *

import tkinter.messagebox
import threading
import webbrowser
from PIL import Image,ImageTk
from playsound import playsound
root=Tk()
root.geometry('1000x650')
root.resizable(0,0)
root.title('alun')
frame_class=[]
def frame_group():
    specification=[(1000,90),(180,530),(820,530),(1000,30)]#容器规格，(标题栏，工具栏，首页,底部状态)
    specification_place=[(-2,0),(0,90),(180,90),(0,620)]
    # color=['#0000FF','#0000CD','#191970' ,'#00008B','#000080','#4169E1','#6495ED','#B0C4DE','#778899','#708090','#1E90FF','red']
    for i in specification:
        width=i[0]
        height=i[1]
        frame=tk.Frame(root, width=width, height=height,)
        frame_class.append(frame)
    for i1,i2 in zip(frame_class,specification_place):
        i1.place(x=i2[0],y=i2[1])
frame_group()
#底部状态栏
bar_title=tk.Label(frame_class[3],text='版本号1.1.0',width=14, height=1)
bar_title.place(x=0,y=0)
bar_title=tk.Label(frame_class[3],text='进度:',width=4, height=1)
bar_title.place(x=653,y=0)
bar_canvas = tk.Canvas(frame_class[3], width=240, height=20, bg="#F5F5F5", borderwidth=0)
bar_canvas.place(x=700,y=0)

#状态栏函数
def change_schedule(now_schedule, all_schedule):
    bar_canvas.coords(fill_rec, (5, 5, 6 + (now_schedule / all_schedule) * 240, 25))
    root.update()
    mini.set(str(round(now_schedule / all_schedule * 100, 2)) + '%')
    if round(now_schedule / all_schedule * 100, 2) == 100.00:
        mini.set("完成")
# 进度条以及完成程度
mini = tk.StringVar()
user_date= tk.StringVar()
out_rec = bar_canvas.create_rectangle(5, 5, 240, 25, outline="blue", width=1)
fill_rec = bar_canvas.create_rectangle(5, 5, 5, 25, outline="", width=0, fill="blue")
tk.Label(frame_class[3], textvariable=mini).place(x=950,y=0)
tk.Label(frame_class[3], textvariable=user_date).place(x=350,y=0)
#11个容器
color=['#0000FF','#0000CD','#191970' ,'#00008B','#000080','#4169E1','#6495ED','#B0C4DE','#778899','#708090','#1E90FF','red']
second_level_class=[]
f_home=frame_class[2]#首页
for i in range(14):
    f=tk.Frame(f_home,width=820,height=560)
    f.place(x=0,y=0)
    second_level_class.append(f)

#子页面控制
def IO(io):
    change_schedule(0, 100)
    for i in second_level_class:
        i.place_forget()
    io.place(x=0,y=0)
    print(i)

def newIO(io):
    change_schedule(0, 100)
    for i in second_level_class:
        i.place_forget()
    io.place(x=0,y=0)
    str_huax()

#top
f_top=frame_class[0]#标题栏
topimg=tk.PhotoImage(file='img/topimg.png')
l_topimg=tk.Label(f_top,image=topimg).place(x=5,y=0)

#tool
f_tool=frame_class[1]
toolimg=tk.PhotoImage(file='img/tool.png')
l_toolimg=tk.Label(f_tool,image=toolimg).place(x=0,y=0)
#工具栏控件
f_tool=frame_class[1]
icon_class=[]
button_class =[]
def set_toolicon():
    icon_path = ['img/icon10.png', 'img/icon9.png', 'img/icon8.png', 'img/icon7.png', 'img/icon6.png', 'img/icon5.png',
                 'img/icon4.png', 'img/icon3.png', 'img/icon2.png', 'img/icon1.png','img/icon0.png','img/icon00.png']#素材地址
    for i in icon_path:
        icon = tk.PhotoImage(file=i)
        icon_class.append(icon)
    button0=tk.Button(f_tool,image =icon_class[0],command=lambda :IO(second_level_class[0])).place(x=5,y=5)
    button1 =tk.Button(f_tool, image=icon_class[1], command=lambda: IO(second_level_class[1])).place(x=5, y=52)
    button2 = tk.Button(f_tool, image=icon_class[2], command=lambda: IO(second_level_class[2])).place(x=5, y=99)
    button3 = tk.Button(f_tool, image=icon_class[3], command=lambda: IO(second_level_class[3])).place(x=5, y=146)
    button4 = tk.Button(f_tool, image=icon_class[4], command=lambda: IO(second_level_class[4])).place(x=5, y=193)
    button5 = tk.Button(f_tool, image=icon_class[5], command=lambda: IO(second_level_class[5])).place(x=5, y=240)
    button6 = tk.Button(f_tool, image=icon_class[6], command=lambda: IO(second_level_class[6])).place(x=5, y=287)
    # button7 = tk.Button(f_tool, image=icon_class[7], command=lambda: IO(second_level_class[7])).place(x=5, y=299)
    button8 = tk.Button(f_tool, image=icon_class[8], command=lambda: IO(second_level_class[8])).place(x=5, y=334)
    button9 = tk.Button(f_tool, image=icon_class[9], command=lambda: IO(second_level_class[9])).place(x=5, y=381)
    button10 = tk.Button(f_tool, image=icon_class[10], command=lambda: IO(second_level_class[10])).place(x=5, y=428)
    button11 = tk.Button(f_tool, image=icon_class[11], command=lambda: newIO(second_level_class[11])).place(x=5, y=475)
set_toolicon()
#写文件
def savepage(page, title):
    mm = askokcancel('提示', '确认下载吗？')
    if mm == True:
        print('开始啦')
        print(page)
        file = Document()
        file.add_heading(str(title), level=2)
        paragraph = file.add_paragraph(str(page))
        paragraph.add_run('来源').bold = True
        file.save(str(title) + '.docx')
        print('成功')
        kk = showinfo('提示', '下载完成')
#11个页面背景图片设置
#['Alipay', 'taobao', 'bilibili', 'baidu', 'tuniu', 'boos', 'netease cloud', 'Meituan', 'xuexingnet', 'jingdong', 'about us']
img_class=[]
def set_secandpage_img():
    page_secand_img = ['img/01.png', 'img/02.png', 'img/03.png', 'img/04.png', 'img/05.png', 'img/06.png', 'img/07.png',
                       'img/08.png', 'img/09.png', 'img/10.png']
    for i in page_secand_img:
        img = tk.PhotoImage(file=i)
        img_class.append(img)
    for i ,page in zip(img_class,second_level_class):
        img_sit = tk.Label(page, image=i,compound = tk.CENTER,borderwidth=0).place(x=0, y=0)
set_secandpage_img()

#小套路
def time2():
    time.sleep(15)
    for i in range(10,83):
        time.sleep(0.1)
        change_schedule(i, 100)
# 详细数据
def opendatas(path):
    webbrowser.open(path)

def date():
    dates=['正在部署爬虫','正在拉起登陆窗口','正在处理登陆信息','爬虫准备中','正在爬虫','正在获取数据标签','正在获取数据']
    for i in dates:
        time.sleep(2)
        user_date.set(i)
        time.sleep(3)
#支付宝：
f_alipay=second_level_class[0]

re_value=[]
alipay=[]
def Alipay():

    def str_Alippay():
        new_Alipay = ZhiFuBao()
        resturd = new_Alipay.go()
        user_data_values = (list(resturd[0].values()))
        print(user_data_values)
        time_data_values = (list(resturd[1].values()))
        print(time_data_values)
        goods_data_values = (list(resturd[2].values()))
        global re_value
        re_value = goods_data_values
        # print(len(goods_data_values),len(re_value))
        # print(re_value)
        var1.set('用户名:' + str(user_data_values[0]))
        var2.set('账单时间:' + str(time_data_values[0]))
        time.sleep(2)
        alipay.append(re_value[0])
        alipay.append(re_value[1])
        alipay.append(re_value[3])
        alipay.append(re_value[14])
        alipay.append(re_value[13])

        for i in range(83, 100):
            time.sleep(0.3)
            change_schedule(i, 100)
        e = 0
        for x in sit_x:
            for y in sit_y:
                laber = tk.Label(f_alipay, justify='left', wraplength=780, text=re_value[e], font=('Arial', 25),
                                 bg='#FFFFFF', fg='#000000').place(x=x, y=y+20)
                e = e + 1



    mm = askokcancel('提示', '操作之前，请打开您的支付宝扫一扫功能，准备识别二维码')
    if mm == True:
        print('开始啦')
        for i in range(10):
            change_schedule(i, 100)
        new_Alipay = threading.Thread(target=str_Alippay)  # target是要执行的函数名（不是函数），args是函数对应的参数，以元组的形式存在
        t1 = threading.Thread(target=date)
        t2 = threading.Thread(target=time2)
        new_Alipay.start()
        t1.start()
        t2.start()
        resturd = new_Alipay
    else:
        print('tuichul')
        pass

var1=tk.StringVar()
var2=tk.StringVar()
var3=tk.StringVar()
var1.set('用户名:'+'____')
var2.set('账单时间:'+'-----')
e_values=['**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**','**,**,**',]
re_keys=['全年交易额:','全年顾客:','全年收入:','全年收入总金额:','全年支出总金额:','全年日均余额:','月均账户收入:','月均账户支出(元):','全年总成交金额(元):','总成交笔数(笔):','全年总退款金额(元):','总退款笔数(笔):','总实收金额(元):','平均笔单价(元):','日均成交金额(元):','日均成交笔数:','全年总顾客数(人):','老顾客总数(人):','新顾客总数(人):','人均消费金额(元):','人均消费笔数(笔):','累计最高消费金额(元):','累计最高消费笔数(笔):',]
var1.set('用户名:'+'____')
var2.set('账单时间:'+'-----')
sit_x=[40,240,440,640]
sit_y=[250,310,370,430,]
a = 0
for x in sit_x:
    for y in sit_y:
        laber = tk.Label(f_alipay, justify='left', wraplength=780, text=re_keys[a], font=('Arial', 11),
                     bg='#FFFFFF', fg='#000000').place(x=x, y=y)
        a=a+1
alipay_user=tk.Label(f_alipay,textvariable=var1,font=('Arial', 25),bg='#FFFFFF',fg='#000000').place(x=40,y=150)
alipay_billtime=tk.Label(f_alipay,textvariable=var2,font=('Arial', 25),bg='#FFFFFF',fg='#000000').place(x=40,y=200)
alipaylogin=tk.PhotoImage(file='img/iconx4.png')
alipay_star=tk.Button(f_alipay,command=Alipay,image=alipaylogin).place(x=400,y=90)
alipaydw=tk.PhotoImage(file='img/iconx5.png')
alipay_download=tk.Button(f_alipay,command=lambda:savepage(re_value,'alipaydata'),image=alipaydw).place(x=550,y=90)
#淘宝
f_taobao=second_level_class[1]
tb_var6=tk.StringVar()
tb_var7=tk.StringVar()
tb_var8=tk.StringVar()
tb_var9=tk.StringVar()
tb_var10=tk.StringVar()
tb_var11=tk.StringVar()
tb_var12=tk.StringVar()
tb_var6.set('咕噜咕噜，让我看看是谁来了')
tb_var7.set('想知道你的消费数据吗')
tb_var8.set('叫我一声小可爱，我就告诉你')
tb_var9.set('嘻嘻嘻点击爬虫 查看详细数据')
tb_var10.set('嘻嘻嘻点击爬虫 查看详细数据')
tb_var11.set('嘻嘻嘻点击爬虫 查看详细数据')
def taoyandedongxi(taobaoclass):
    tb_var6.set(taobaoclass[0][0])

    if taobaoclass[1] == []:
        tb_var7.set('没有查询到,下次试试吧')
    else:
        tb_var7.set(str(taobaoclass[1]))
    tb_var8.set(str(taobaoclass[2]))
    tb_var9.set(str(taobaoclass[3]))
    tb_var10.set(str(taobaoclass[4]))
    tb_var10.set(str(taobaoclass[5]))
    for i in range(82,100):
        change_schedule(i, 100)

def taobaojiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)

def taobaotime():
    print('开始啦')
    for i in range(10,82):
        change_schedule(i, 100)
taobaodw=[]
def taobaostr():
    def taogo():
        mm = askokcancel('提示', '操作之前，请打开您的taobao扫一扫功能，准备识别二维码')
        if mm == True:
            print('开始啦')
            for i in range(10):
                change_schedule(i, 100)
            str_taobao()
        else:
            print('tuichul')
            pass
    def str_taobao():
        result = taobao_go()
        global taobaodw
        taobaodw=result
        taoyandedongxi(result)


    new_taobao = threading.Thread(target=taogo)
    new_taobao.start()
    t1 = threading.Thread(target=taobaojiashuju)
    t2 = threading.Thread(target=time3)
    t1.start()
    t2.start()






taoblogin=tk.PhotoImage(file='img/iconx12.png')
taobmore=tk.PhotoImage(file='img/iconx13.png')
taobao_star=tk.Button(f_taobao,command=taobaostr,image=taoblogin).place(x=470,y=390)
taobao_xiangxishuju=tk.Button(f_taobao,command=lambda :savepage(taobaodw,'taobao'),image=taobmore).place(x=600,y=390)
taobao_1=tk.Label(f_taobao,textvariable=tb_var6,font=('Arial', 10),bg='#FFFFFF',fg='#000000').place(x=60,y=200)
taobao_2=tk.Label(f_taobao,textvariable=tb_var7,font=('Arial', 10),bg='#FFFFFF',fg='#000000').place(x=60,y=250)
taobao_3=tk.Label(f_taobao,textvariable=tb_var8,font=('Arial', 10),bg='#FFFFFF',fg='#000000').place(x=60,y=300)
taobao_4=tk.Label(f_taobao,textvariable=tb_var9,font=('Arial', 10),bg='#FFFFFF',fg='#000000').place(x=60,y=350)
taobao_5=tk.Label(f_taobao,textvariable=tb_var10,font=('Arial', 10),bg='#FFFFFF',fg='#000000').place(x=60,y=399)
taobao_6=tk.Label(f_taobao,textvariable=tb_var11,font=('Arial', 10),bg='#FFFFFF',fg='#000000').place(x=60,y=459)

#bilibili
f_bilibili=second_level_class[2]
bb_var1=tk.StringVar()
bb_var2=tk.StringVar()
bb_var3=tk.StringVar()
bb_var4=tk.StringVar()
bb_var5=tk.StringVar()
bb_var1.set('欢迎来到哔哩看数据')
bb_var2.set('想知道你的爱好嘛')
bb_var3.set('要不要了解一下自己的浏览数据？')
bb_var4.set('兴趣爱好？')
bb_var5.set('点击开始，尝试一下！')

def bilibilixiangxishuju(resturt):
    bb_var1.set(str(resturt[3]))
    bb_var2.set(str(resturt[4]))
    bb_var3.set(str(resturt[0]))
    bb_var4.set(str(resturt[1]))
    bb_var5.set(str(resturt[2]))
def bilibilijiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)
def time3():
    time.sleep(15)
    for i in range(10,83):
        change_schedule(i, 100)
bilibili_resurt=[]
bbdw=[]
def bilibiliw():
    def bil():
        mm = askokcancel('提示', '操作之前，请打开您的bilibil扫一扫功能，准备识别二维码')
        if mm == True:
            print('开始啦')
            for i in range(10):
                change_schedule(i, 100)
            str_bilibili()
        else:
            print('tuichul')
            pass
    def str_bilibili():
        new_bilibilis=bilibili_go()
        bilibilixiangxishuju(new_bilibilis)
        global bbdw
        bbdw=new_bilibilis


    new_bilibili = threading.Thread(target=bil)
    new_bilibili.start()
    t1 = threading.Thread(target=bilibilijiashuju)
    t2 = threading.Thread(target=time3)
    t1.start()
    t2.start()

bilibillogin=tk.PhotoImage(file='img/iconx2.png')
bilibilmore=tk.PhotoImage(file='img/iconx3.png')
bilibil_star=tk.Button(f_bilibili,command=bilibiliw,image=bilibillogin).place(x=70,y=420)
bilibil_xiangxishuju=tk.Button(f_bilibili,command=lambda :savepage(bbdw,'bilibili'),image=bilibilmore).place(x=210,y=420)
bilibili_1=tk.Label(f_bilibili,textvariable=bb_var1,font=('Arial',10),bg='#FFFFFF',fg='#000000').place(x=275,y=220)
bilibili_2=tk.Label(f_bilibili,textvariable=bb_var2,font=('Arial',10),bg='#FFFFFF',fg='#000000').place(x=275,y=255)
bilibili_3=tk.Label(f_bilibili,textvariable=bb_var3,font=('Arial',10),bg='#FFFFFF',fg='#000000').place(x=275,y=303)
bilibili_4=tk.Label(f_bilibili,textvariable=bb_var4,font=('Arial',10),bg='#FFFFFF',fg='#000000').place(x=275,y=342)
bilibili_5=tk.Label(f_bilibili,textvariable=bb_var5,font=('Arial',10),bg='#FFFFFF',fg='#000000').place(x=275,y=385)

#baidu
f_baidu=second_level_class[3]
bdyx_var1=tk.StringVar()
bdyx_var2=tk.StringVar()
bdyx_var3=tk.StringVar()
bdyx_var4=tk.StringVar()
bdyx_var5=tk.StringVar()
bdyx_var1.set('今天天气很不错')
bdyx_var2.set('今天有百度吗，看看都浏览了哪些信息？')
bdyx_var3.set('最热关键词，每天的热点，你都知道吗')
bdyx_var4.set('点击开始，给你送上专属称号')
bdyx_var5.set('快来试试吧')
def baiduxiangxishuju(er):
    bdyx_var1.set(er[3])
    bdyx_var2.set(er[2][0])
    bdyx_var3.set(er[1][0:3])
    bdyx_var4.set(str(er[0]))
    # bdyx_var5.set(er[4][0])
def baidujiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)
def time4():
    time.sleep(15)
    for i in range(10,83):
        time.sleep(0.7)
        change_schedule(i, 100)
bddw=[]
def baiduww():
    mm = askokcancel('提示', '操作之前，请打开您的baidu扫一扫功能，准备识别二维码')
    if mm == True:
        print('开始啦')
        for i in range(10):
            change_schedule(i, 100)
        baiduw()
    else:
        print('tuichul')
        pass
baiduresturd=[]
def baiduw():
    def str_baidu():
        global bddw
        new_baidus=baidu_go()
        bddw=new_baidus
        baiduxiangxishuju(new_baidus)

    new_baidu = threading.Thread(target=str_baidu)
    new_baidu.start()
    t1 = threading.Thread(target=baidujiashuju)
    t2 = threading.Thread(target=time4)
    t1.start()
    t2.start()


baidlogin=tk.PhotoImage(file='img/iconx8.png')
baidumore=tk.PhotoImage(file='img/iconx9.png')
baidu_star=tk.Button(f_baidu,command=baiduww,image=baidlogin).place(x=450,y=460)
baidu_xiangxishuju=tk.Button(f_baidu,command=lambda :savepage(bddw,'baidu.txt'),image=baidumore).place(x=590,y=460)
baidu_1=tk.Label(f_baidu,textvariable=bdyx_var1,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=62,y=235)
baidu_2=tk.Label(f_baidu,textvariable=bdyx_var2,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=62,y=290)
baidu_3=tk.Label(f_baidu,textvariable=bdyx_var3,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=62,y=342)
baidu_4=tk.Label(f_baidu,textvariable=bdyx_var4,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=62,y=400)
baidu_5=tk.Label(f_baidu,textvariable=bdyx_var5,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=62,y=460)


#tuniu
f_tuniu=second_level_class[4]
tuniu_var1=tk.StringVar()
tuniu_var2=tk.StringVar()
tuniu_var3=tk.StringVar()
tuniu_var4=tk.StringVar()
tuniu_var5=tk.StringVar()
tuniu_var6=tk.StringVar()
tuniu_var7=tk.StringVar()
tuniu_var1.set('春风吹战鼓擂，新的一年往哪颓')
tuniu_var2.set('今天雨好大，好像不适合出去玩呢')
tuniu_var3.set('想知道你一年都去了哪些地方吗')
tuniu_var4.set('想知道偌大的中国都有哪些地方还没涉足吗')
tuniu_var5.set('点击这里 查看详细数据')
tuniu_var6.set('点击这里 查看详细数据')
tuniu_var7.set('点击这里 查看详细数据')
def tuniuxiangxishuju(tuniuclass):
    tuniu_var1.set(tuniuclass[0])
    tuniu_var2.set(tuniuclass[1])
    tuniu_var3.set(tuniuclass[2])
    tuniu_var4.set(tuniuclass[3])
    tuniu_var5.set(tuniuclass[4])
    tuniu_var6.set(tuniuclass[5])
    tuniu_var7.set(tuniuclass[6])
def tuniujiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)
def time5():
    time.sleep(15)
    for i in range(10,83):
        time.sleep(0.7)
        change_schedule(i, 100)
tuniudw=[]
def tuniuw():
    def str_tuniu():
        global tuniudw
        new_tunius=tuniugo()
        tuniudw=new_tunius
        tuniuxiangxishuju(new_tunius)

    new_tuniu = threading.Thread(target=str_tuniu)
    new_tuniu.start()
    t1 = threading.Thread(target=tuniujiashuju)
    t2 = threading.Thread(target=time5)

    t1.start()
    t2.start()
tu1=tk.PhotoImage(file='img/iconx18.png')
tu2=tk.PhotoImage(file='img/iconx19.png')
tuniu_star=tk.Button(f_tuniu,command=tuniuw,image=tu1).place(x=350,y=100)
tuniu_xiangxishuju=tk.Button(f_tuniu,command=lambda :savepage(tuniudw,'tuniu'),image=tu2).place(x=550,y=100)
tuniu_1=tk.Label(f_tuniu,textvariable=tuniu_var1,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=55,y=240)
tuniu_2=tk.Label(f_tuniu,textvariable=tuniu_var2,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=55,y=300)
tuniu_3=tk.Label(f_tuniu,textvariable=tuniu_var3,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=55,y=350)
tuniu_4=tk.Label(f_tuniu,textvariable=tuniu_var4,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=55,y=410)
tuniu_5=tk.Label(f_tuniu,textvariable=tuniu_var5,font=('Arial', 12),bg='#FFFFFF',fg='#000000').place(x=55,y=465)

#qqemail
f_qqemail=second_level_class[5]
qqemail_var1=tk.StringVar()
qqemail_var2=tk.StringVar()
qqemail_var3=tk.StringVar()
qqemail_var4=tk.StringVar()
qqemail_var5=tk.StringVar()
qqemail_var1.set('2020 小B还没有找到工作，你呢')
qqemail_var2.set('我猜你是个程序员')
qqemail_var3.set('正在点燃箱子里的火')
qqemail_var4.set('正在活动指关节')
qqemail_var5.set('点击这里 查看详细数据')
def qqemailxiangxishuju(er):
    qqemail_var1.set(er[0])
    qqemail_var2.set(er[1])
    qqemail_var3.set(er[2])
    # qqemail_var4.set(er[2][1])
    # qqemail_var5.set(er[2][])
def qqemailjiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)
def time6():
    time.sleep(15)
    for i in range(10,83):
        time.sleep(0.7)
        change_schedule(i, 100)
resturds=[]
qqdata=[]
qqdw=[]
def qqemailw():
    def str_qqemail():
        new_qqemails=emailgo()
        global qqdw
        qqdw=new_qqemails
        qqdata.append(new_qqemails[-1])
        qqemailxiangxishuju(new_qqemails)

    new_qqemail = threading.Thread(target=str_qqemail)
    new_qqemail.start()
    t1 = threading.Thread(target=qqemailjiashuju)
    t2 = threading.Thread(target=time6)
    t1.start()
    t2.start()
qq1=tk.PhotoImage(file='img/iconx14.png')
qq2=tk.PhotoImage(file='img/iconx15.png')
qqemail_star=tk.Button(f_qqemail,command=qqemailw,image=qq1).place(x=480,y=430)
qqemail_xiangxishuju=tk.Button(f_qqemail,command=lambda :savepage(qqdw,'mail'),image=qq2).place(x=630,y=430)
qqemail_1=tk.Label(f_qqemail,textvariable=qqemail_var1,font=('Arial', 15),bg='#FFFFFF',fg='#000000').place(x=60,y=220)
qqemail_2=tk.Label(f_qqemail,textvariable=qqemail_var2,font=('Arial', 15),bg='#FFFFFF',fg='#000000').place(x=60,y=300)
qqemail_3=tk.Label(f_qqemail,textvariable=qqemail_var3,font=('Arial', 15),bg='#FFFFFF',fg='#000000').place(x=60,y=370)



#wangyiyun
f_wangyiyun=second_level_class[6]
wangyiyun_var1=tk.StringVar()
wangyiyun_var2=tk.StringVar()



wangyiyun_var1.set('欢迎，陌上花开')
wangyiyun_var2.set('好久不见，今天想听点什么')



def oppppp():
    playsound('1.mp3')
def musicopen(path):
    musicop = threading.Thread(target=oppppp)
    musicop.start()

def wangyiyunxiangxishuju(a):
    wangyiyun_var1.set(a[1])
    wangyiyun_var2.set(a[0])

def wangyiyunkais():
    mm = askokcancel('提示', '操作之前，请打开您的网易云扫一扫功能，准备识别二维码')
    if mm==True:
        print('开始啦')
        wangyiyunw()
    else:
        print('tuichul')
        pass
def wangyiyunjiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)
def time7():
    time.sleep(15)
    for i in range(10,83):
        time.sleep(0.7)
        change_schedule(i, 100)
wwwww=[]
wydw=[]
def wangyiyunw():
    def str_wangyiyun():
        global wydw
        new_wangyiyuns=go()
        wydw=new_wangyiyuns
        wangyiyunxiangxishuju(wydw)

    str_wangyiyun()
wangyiyun_open=tk.Button(f_wangyiyun,command=lambda :musicopen('1.mp3'),text='播放').place(x=500,y=100)
wangyiyun_star=tk.Button(f_wangyiyun,command=wangyiyunkais,text='登陆').place(x=500,y=470)
wangyiyun_xiangxishuju=tk.Button(f_wangyiyun,command=lambda :savepage(wydw,'wangyiyun'),text='下载').place(x=600,y=470)
wangyiyun_1=tk.Label(f_wangyiyun,textvariable=wangyiyun_var1,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=40,y=250)
wangyiyun_2=tk.Label(f_wangyiyun,textvariable=wangyiyun_var2,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=40,y=395)


#xuexinw
f_xuexinw=second_level_class[8]
xx_var1=tk.StringVar()
xx_var2=tk.StringVar()
xx_var3=tk.StringVar()
xx_var4=tk.StringVar()
xx_var5=tk.StringVar()
xx_var6=tk.StringVar()
xx_var7=tk.StringVar()
xx_var8=tk.StringVar()
xx_var9=tk.StringVar()
xx_var10=tk.StringVar()
xx_var11=tk.StringVar()
xx_var12=tk.StringVar()


im1 = Image.open('user.png') #支持相对或绝对路径，支持多种格式
im2=ImageTk.PhotoImage(im1)
xuex=[]
xxdw=[]
def xuexinww():
    def str_xuexinw():
        user_date.set('爬虫软件开始')
        for i in range(10):
            change_schedule(i, 100)
        name=xuexinw_e1.get()
        password=xuexinw_e2.get()
        new_xuexinws=Login(name,password)
        resturt=new_xuexinws.xx()
        xuex.append(resturt)
        global xxdw
        xxdw=resturt
        print(resturt)
        for i in range(0, 100):
            change_schedule(i, 100)

        canvas = tk.Canvas(f_xuexinw, width=200, height=267)
        global im1
        global im2
        im1 = Image.open('user_info.jpg')  # 支持相对或绝对路径，支持多种格式
        im2 = ImageTk.PhotoImage(im1)
        canvas.create_image(0,0,anchor='nw',image=im2)
        canvas.place(x=550, y=110)
        xx_var1.set(resturt[0])
        global huaxiang_name
        huaxiang_name= (resturt[0])
        xx_var2.set(resturt[1])
        global huaxiang_age
        huaxiang_age = (resturt[2])
        global huaxiang_six
        huaxiang_six = (resturt[1])
        global huaxiang_xueli
        huaxiang_xueli = (resturt[6])
        xx_var3.set(resturt[2])
        xx_var4.set(resturt[3])
        xx_var5.set(resturt[4])
        xx_var6.set(resturt[5])
        xx_var7.set(resturt[6])
        xx_var8.set(resturt[7])
        xx_var9.set(resturt[8])
        xx_var10.set(resturt[9])
        xx_var11.set(resturt[10])
        xx_var12.set(resturt[11])
    str_xuexinw()

xuexinw_l1=tk.Label(f_xuexinw,text='用户名').place(x=500,y=400)
xuexinw_l2=tk.Label(f_xuexinw,text='密码').place(x=500,y=430)
xuexinw_e1=tk.Entry(f_xuexinw,show=None)
xuexinw_e1.place(x=550,y=400)
xuexinw_e2=tk.Entry(f_xuexinw,show='*')
xuexinw_e2.place(x=550,y=430)
xuexinw_star=tk.Button(f_xuexinw,command=xuexinww,text='查询').place(x=560,y=470)
xuexinw_xiangxishuju=tk.Button(f_xuexinw,command=lambda :savepage(xxdw,'xuexinw'),text='下载').place(x=660,y=470)
xuexinw_1=tk.Label(f_xuexinw,textvariable=xx_var1,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=70,y=190)
xuexinw_2=tk.Label(f_xuexinw,textvariable=xx_var2,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=70,y=242)
xuexinw_3=tk.Label(f_xuexinw,textvariable=xx_var3,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=70,y=295)
xuexinw_4=tk.Label(f_xuexinw,textvariable=xx_var4,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=70,y=350)
xuexinw_5=tk.Label(f_xuexinw,textvariable=xx_var5,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=70,y=399)
xuexinw_6=tk.Label(f_xuexinw,textvariable=xx_var6,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=70,y=453)
xuexinw_7=tk.Label(f_xuexinw,textvariable=xx_var7,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=370,y=190)
xuexinw_8=tk.Label(f_xuexinw,textvariable=xx_var8,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=370,y=242)
xuexinw_9=tk.Label(f_xuexinw,textvariable=xx_var9,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=370,y=295)
xuexinw_10=tk.Label(f_xuexinw,textvariable=xx_var10,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=370,y=350)
xuexinw_11=tk.Label(f_xuexinw,textvariable=xx_var11,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=370,y=399)
xuexinw_12=tk.Label(f_xuexinw,textvariable=xx_var12,font=('Arial', 13),bg='#FFFFFF',fg='#000000').place(x=370,y=453)


#jingdong
f_jingdong=second_level_class[9]
jd_var1=tk.StringVar()
jd_var2=tk.StringVar()
jd_var3=tk.StringVar()
jd_var4=tk.StringVar()
jd_var5=tk.StringVar()
jd_var6=tk.StringVar()
jd_var7=tk.StringVar()
jd_var1.set('很高兴遇到你')
jd_var2.set('2020年7月5号是你最近一次登陆淘宝，有找到喜欢的东西吗')
jd_var3.set('2019年你共浏览商品13632次，下单120次，获得称号剁手狂魔')
jd_var4.set('2019年6月 你给***购买了情人节礼物，她一定很幸福吧')
jd_var5.set('点击这里 查看详细数据')
jd_var6.set('点击这里 查看详细数据')
jd_var7.set('点击这里 查看详细数据')
def jingdongxiangxishuju(jingdclass):
    jd_var1.set(jingdclass[0])
    jd_var2.set(jingdclass[1])
    jd_var3.set(jingdclass[2])
    jd_var4.set(jingdclass[3])
    jd_var5.set(jingdclass[4])
    jd_var6.set(jingdclass[5])
    jd_var7.set(jingdclass[6])
def jingdongjiashuju():
    user_date.set('爬虫软件开始')
    time.sleep(8)
    dates = ['正在部署爬虫', '正在拉起登陆窗口', '正在处理登陆信息', '爬虫准备中', '正在爬虫', '正在获取数据标签', '正在获取数据','个人信息获取成功',
             '成功提取到地址信息','购物车数据已获取','浏览历史已获取','订单信息已获取','正在完善个人数据','正在提取标签','数据清洗中','数据准备完毕']
    for i in dates:
        user_date.set(i)
        time.sleep(6)
def time10():
    time.sleep(15)
    for i in range(10,83):
        time.sleep(0.7)
        change_schedule(i, 100)
jddw=[]
def jingdongw():
    def str_jingdong():
        new_jingdongs=jdgo()
        global jddw
        jddw=new_jingdongs
        x=jingdongxiangxishuju(new_jingdongs)
    for i in range(10):
        time.sleep(0.1)
        change_schedule(i, 100)
    new_jingdong = threading.Thread(target=str_jingdong)
    new_jingdong.start()
    t1 = threading.Thread(target=jingdongjiashuju)
    t2 = threading.Thread(target=time10)

    t1.start()
    t2.start()
jd1=tk.PhotoImage(file='img/iconx6.png')
jd2=tk.PhotoImage(file='img/iconx7.png')
jingdong_star=tk.Button(f_jingdong,command=jingdongw,image=jd1).place(x=480,y=420)
jingdong_xiangxishuju=tk.Button(f_jingdong,command=lambda :savepage(jddw,'jingdong.txt'),image=jd2).place(x=630,y=420)
jingdong_1=tk.Label(f_jingdong,textvariable=jd_var1,font=('Arial', 11),bg='#FFFFFF',fg='#000000').place(x=60,y=233)
jingdong_2=tk.Label(f_jingdong,textvariable=jd_var2,font=('Arial', 11),bg='#FFFFFF',fg='#000000').place(x=60,y=280)
jingdong_3=tk.Label(f_jingdong,textvariable=jd_var3,font=('Arial', 11),bg='#FFFFFF',fg='#000000').place(x=60,y=326)
jingdong_4=tk.Label(f_jingdong,textvariable=jd_var4,font=('Arial', 11),bg='#FFFFFF',fg='#000000').place(x=60,y=373)
jingdong_5=tk.Label(f_jingdong,textvariable=jd_var5,font=('Arial', 11),bg='#FFFFFF',fg='#000000').place(x=60,y=422)


#关于我们
f_guany=second_level_class[10]
aboutimg=tk.PhotoImage(file='img/11.png')
guanyu_img_sit=tk.Label(f_guany,image=aboutimg).place(x=0,y=0)
github=tk.Label(f_guany,text='https://github.com/sdhushu/VIKE-').place(x=40,y=300)
qqemailabout=tk.Label(f_guany,text='QQ群ID :715433204').place(x=40,y=330)
sentens=tk.Label(f_guany,text='诚挚邀请各位一同参与艾伦项目维护').place(x=40,y=360)
#人物画像
f_huax=second_level_class[11]
img=Image.open('01x.png')
im3=ImageTk.PhotoImage(img)
xiaofeishuju = tk.PhotoImage(file='01x.png')
sumiao = tk.PhotoImage(file='data_info.jpg')
huaxiang_name='None'
huaxiang_age='None'
huaxiang_six='None'
huaxiang_xueli='None'
def str_huax():
    print(qqdata[0][0])
    cloudcc(str(qqdata[0][0]))
    ResizeImage('标签云效果图.png','reimg.png',260,130)
    # vsgo([alipaydata[0])
    # ResizeImage('xxx.png', 'kshimg.png',320,240)
    global xiaofeishuju
    xiaofeishuju = tk.PhotoImage(file='reimg.png')
    global keshihuashuju
    keshihuashuju = tk.PhotoImage(file='kshimg.png')
    global ditushuju
    ResizeImage('adress.png', 'newadress.png',288,243)
    global sumiao
    ResizeImage('user_info1.png', 'user_info2.png', 115, 153)
    sumiao = tk.PhotoImage(file='user_info2.png')
    ditushuju = tk.PhotoImage(file='newadress.png')
    sumiaoshuju_sit=tk.Label(f_huax, image=sumiao).place(x=350, y=150)
    xiaofeishuju_sit = tk.Label(f_huax, image=xiaofeishuju).place(x=50, y=50)
    keshihuashuju_sit = tk.Label(f_huax, image=keshihuashuju).place(x=475, y=0)
    ditushuju_sit = tk.Label(f_huax, image=ditushuju).place(x=485, y=260)
    w1 = tk.Label(f_huax, text='人际关系云图',font=('Arial', 9)).place(x=140, y=210)
    w2 = tk.Label(f_huax, text='收入支出图',font=('Arial', 9)).place(x=605, y=230)
    w3 = tk.Label(f_huax, text='性格测试',font=('Arial', 9)).place(x=140, y=490)
    w4 = tk.Label(f_huax, text='所在商圈位置',font=('Arial', 9)).place(x=605, y=260)
    w5 = tk.Label(f_huax, text='姓名',font=('Arial', 9)).place(x=350, y=300)
    w6 = tk.Label(f_huax, text='年龄',font=('Arial', 9)).place(x=350, y=325)
    w7 = tk.Label(f_huax, text='性别',font=('Arial', 9)).place(x=350, y=350)
    w8 = tk.Label(f_huax, text='学历',font=('Arial', 9)).place(x=350, y=375)
    w9 = tk.Label(f_huax, text=huaxiang_name,font=('Arial', 9)).place(x=380,  y=300)
    w10 = tk.Label(f_huax, text=huaxiang_age,font=('Arial', 9)).place(x=380, y=325)
    w11 = tk.Label(f_huax, text=huaxiang_six,font=('Arial', 9)).place(x=380, y=350)
    w12 = tk.Label(f_huax, text=huaxiang_xueli,font=('Arial', 9)).place(x=380, y=375)
    a = random.randint(0, 8)
    list_hobby = [
        '你需要别人喜欢你和欣赏你，但你通常对自己要求苛刻。\n虽然你在个性上有一些弱点，但你通常能够设法加以弥补。\n你在某些方面的能力并没有得到充分发挥，所以还未能变成你的优势。\n从外表来看，你是一个讲求自律和自制的人，但内心却常常焦虑不安。\n有时候，你会强烈怀疑自己是不是做出了正确的决定或正确的事情。\n你倾向于让自己的生活有所改变和变得丰富多彩，在遇到约束和限制时你\n会感到不满。你很自豪自己是一个能够独立思考的人，\n如果没有令人满意的证据，你不会接受别人的观点和说法。\n不过，你也觉得在别人面前过于直言不讳并不是明智之举，\n有时候你很外向，比较容易亲近，也乐于与人交往，但有时候你却很内向，\n比较小心谨慎，而且沉默寡言。你有很多梦想，\n其中有一些看起来相当不实际。',
        '你是典型的本质消极但外表热情的人，好像总是有所追求。\n你追求一切让自己的生活过得更好的事物，却又没办法去正视\n自己天生的悲观和灰暗，所以你会表现得既浮躁急切又愤世嫉俗',
        '你的完美总是表现得过于夸张，有时已经完全脱离真实的生活，\n上升到外太空、异世界层面。所以你的一生与其说在遐想的\n完美中度过，不如说在一种绝对的科幻或玄幻状态中度过，\n你所期望的一切估计只有在科幻或玄幻小说里才能找到并完成吧&oq=你的\n完美总是表现得过于夸张，有时已经完全脱离真实的生活，上升到外太空、\n异世界层面。所以你的一生与其说在遐想的完美中度过，不如说在一\n种绝对的科幻或玄幻状态中度过，你所期望的一切估计\n只有在科幻或玄幻小说里才能找到并完成吧',
        '你的想法天马行空，不着边际，好奇心和观察力总是天花乱坠，\n无人能及。别人根本不可能知道你的下一步是什么，下一个\n问题的答案在哪里。甚至连你自己都不知道做一件事的目的何在，\n意义在哪。你的人生不是在不停的悬疑中进行，就是在荒诞中落幕',
        '虽然你十分相信自己，但是当你面对陌生的环境时，会变得不够自信。\n但是你会承认这种情况，因为你有一个坚强的性格。\n你并不需要投入大量的精力，一般来说你会争取自己想要的',
        '你是很懂得保护自己和观察自己的人，你是向上的积极的，\n你已经给自己制定好了生活目标，你朝着这样的目标在前进着。\n抑郁要找上你也是比较困难的，你很难会抑郁，你面对生活的时候\n不会让自己陷入纠结，你很果断，知道该有什么样的选择，日子过得\n相对轻松',
        '你经常会有迷茫的感觉，你的抑郁几率是很高的，\n你的生活看起来并没有那么的顺利。你其实是很害怕压力的人，面对生活的\n时候你经常会质疑自己。你想要的东西越多，那么你的情绪就变得\n越是不稳定，这样的日子对你来说是很艰难的，你要注意自己的情\n绪。'
        '性情温良，重友谊，性格塌实稳重，但有时也比较狡黠。\n事业心一般，对本职工作能认真对待，但对自己专业以外事物没有太大兴趣\n喜欢有规律的工作和生活，不喜欢冒险，家庭观念强，比较善于理财。',
        '爱幻想，思维较感性，以是否与自己投缘为标准来选择朋友。\n性格显得较孤傲，有时较急躁，有时优柔寡断。事业心较强，\n喜欢有创造性的工作，不喜欢按常规办事。性格倔强，言语犀利，\n不善于妥协。崇尚浪漫的爱情，但想法往往不切合实际。金钱欲望一般\n。',
        '意志力强，头脑冷静，有较强的领导欲，事业心强，不达目的不罢休。\n外表和善，内心自傲，对有利于自己的人际关系比较看重，\n有时显得性格急噪，咄咄逼人，得理不饶人，不利于自己时顽强\n抗争，不轻易认输。思维理性，对爱情和婚姻的看法很现实，对金钱\n的欲望一般。'
    ]
    print(a)
    text= (list_hobby[a])
    xingeceshi = tk.Label(f_huax, text=text, font=('Arial', 7), bg='#FFFFFF', fg='#000000').place(x=20, y=332)

#首页
shouyeimg=tk.PhotoImage(file='img/maddol.png')
shouye_img_sit=tk.Label(second_level_class[13],image=shouyeimg).place(x=0,y=0)
guanyuwomen_img_sit=tk.Label(second_level_class[13],image=shouyeimg).place(x=0,y=0)

root.mainloop()